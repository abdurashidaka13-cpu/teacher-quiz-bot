import datetime
import io
import string
import random
from typing import List
from aiogram import Router, F, Bot
from aiogram.fsm.context import FSMContext
from aiogram.types import Message, CallbackQuery, BufferedInputFile, InlineKeyboardMarkup, InlineKeyboardButton
from openpyxl.styles import Font
from sqlalchemy import select, func, delete

from config import ADMIN_IDS, TIMEZONE, local_now
from database import (
    async_session,
    User,
    Subscription,
    Quiz,
    Question,
    Student,
    StudentAttempt,
    StudentAnswer,
    SystemSettings,
)
from keyboards.bot_keyboards import (
    get_welcome_kb,
    get_teacher_menu,
    get_guide_download_kb,
    get_student_menu,
    get_cancel_kb,
    get_admin_menu,
)
from states import AuthStates, TeacherStates, StudentStates
from utils.docx_parser import parse_docx_questions
from utils.excel_handler import parse_students_excel, generate_results_excel

router = Router(name="teacher")


# FSM blokirovkasini (deadlock) oldini oluvchi middleware
@router.message.middleware()
async def teacher_fsm_cancel_middleware(handler, event: Message, data: dict):
    state: FSMContext = data.get("state")
    current_state = await state.get_state() if state else None
    if current_state is not None and event.text in [
        "Yangi test yaratish",
        "Mening testlarim",
        "Mening obunam",
        "Qo'llanma",
        "Adminga murojaat",
        "Talaba rejimiga o'tish",
        "/start"
    ]:
        await state.clear()
        from handlers.common import cmd_start
        await cmd_start(event, state)
        return
    return await handler(event, data)


# ==========================================
# 1. MODERATOR RO'YXATDAN O'TISH
# ==========================================
@router.callback_query(F.data == "auth_teacher")
async def cb_teacher_reg_start(callback: CallbackQuery, state: FSMContext):
    """Moderatorlikka ro'yxatdan o'tishni boshlash (Ariza jarayoni)"""
    user_id = callback.from_user.id

    async with async_session() as session:
        # Avval ro'yxatdan o'tganligini tekshirish
        stmt = select(User).where(User.id == user_id, User.role == "moderator")
        res = await session.execute(stmt)
        teacher = res.scalar_one_or_none()

        if teacher:
            await callback.message.answer(
                "Siz allaqachon moderator sifatida ro'yxatdan o'tgansiz!",
                reply_markup=get_teacher_menu(is_demo=False),
            )
            await callback.answer()
            return

    await state.set_state(AuthStates.waiting_for_moderator_name)
    await callback.message.delete()  # Eski inline menyuni o'chiramiz
    await callback.message.answer(
        "👨‍🏫 **Moderatorlikka ro'yxatdan o'tish**\n\nIltimos, to'liq ism-familiyangizni kiriting:",
        reply_markup=get_cancel_kb(),
        parse_mode="Markdown",
    )
    await callback.answer()


@router.message(AuthStates.waiting_for_moderator_name)
async def process_teacher_name(message: Message, state: FSMContext):
    """Moderator ismini saqlash va ish joyini so'rash"""
    name = message.text.strip()
    if name in ["Tizimdan chiqish", "Bekor qilish", "/start"]:
        await state.clear()
        from handlers.common import cmd_start
        await cmd_start(message, state)
        return

    if not name or len(name) < 3:
        await message.answer("Iltimos, ismingizni to'liqroq kiriting:")
        return

    await state.update_data(full_name=name)
    await state.set_state(AuthStates.waiting_for_moderator_org)
    await message.answer(
        "Tashkilot nomi (maktabingiz, o'quv markazingiz yoki universitetingiz)ni yozing:"
    )


@router.message(AuthStates.waiting_for_moderator_org)
async def process_teacher_org(message: Message, state: FSMContext, bot: Bot):
    """Ro'yxatdan o'tishni yakunlash, bepul Demo tarifini yoqish va bosh adminlarni xabardor qilish"""
    org = message.text.strip()
    if org in ["Tizimdan chiqish", "Bekor qilish", "/start"]:
        await state.clear()
        from handlers.common import cmd_start
        await cmd_start(message, state)
        return

    if not org:
        await message.answer("Iltimos, tashkilot nomini kiriting:")
        return

    data = await state.get_data()
    full_name = data["full_name"]
    user_id = message.from_user.id
    username = message.from_user.username

    async with async_session() as session:
        # Avval foydalanuvchi mavjudligini tekshirish (masalan, u admin bo'lishi mumkin)
        user_stmt = select(User).where(User.id == user_id)
        existing_user = (await session.execute(user_stmt)).scalar_one_or_none()
        
        if existing_user:
            existing_user.full_name = full_name
            existing_user.username = username
            if existing_user.role != "admin":
                existing_user.role = "moderator"
        else:
            # Yangi foydalanuvchi yaratish
            new_teacher = User(
                id=user_id,
                full_name=full_name,
                username=username,
                role="moderator",
            )
            session.add(new_teacher)

        # Avtomatik bepul Demo tarifini yoqish (1 ta test limiti bilan)
        demo_sub = Subscription(
            user_id=user_id,
            type="free_demo",
            credits=1,
            expires_at=local_now() + datetime.timedelta(days=30),
        )
        session.add(demo_sub)
        await session.commit()

    await state.clear()
    await message.answer(
        f"🎉 **Tabriklaymiz, {full_name}!**\nTizimdan muvaffaqiyatli ro'yxatdan o'tdingiz.\n\n"
        "Sizga bepul **Demo tarif** berildi (1 ta test yaratish limiti bilan).",
        reply_markup=get_teacher_menu(is_demo=True),
        parse_mode="Markdown",
    )

    # Super Adminlarni xabardor qilish
    admin_msg = (
        f"🔔 **YANGI MODERATOR RO'YXATDAN O'TDI**\n\n"
        f"👤 **Ismi:** {full_name}\n"
        f"🆔 **ID:** `{user_id}`\n"
        f"🔗 **Username:** @{username if username else 'yoq'}\n"
        f"🏫 **Tashkilot:** {org}\n"
        f"🎁 **Tarif:** Bepul Demo"
    )
    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(admin_id, admin_msg, parse_mode="Markdown")
        except Exception:
            pass


# ==========================================
# 2. SHABLONLAR# 1. QO'LLANMA (GUIDE)
# ==========================================
@router.message(F.text == "Qo'llanma")
async def show_guide(message: Message):
    """Qisqacha matnli qo'llanmani ko'rsatish va shablon yuklash tugmalari"""
    guide_text = (
        "📖 <b>TEST YARATISH BO'YICHA QISQA QO'LLANMA</b>\n\n"
        "<b>1. Test savollarini yuklash:</b>\n"
        "Word shablonini yuklab oling va savollarni jadvalga to'ldiring. Har bir savol uchun to'g'ri javob va 3 ta noto'g'ri variant to'liq yozilishi shart.\n"
        "<i>Izoh: savol matniga formulalar va chizmalarni rasm shaklida joylashingiz mumkin.</i>\n\n"
        "<b>2. O'quvchilarga login-parol yaratish:</b>\n"
        "Excel shablonini yuklab oling, o'quvchilaringiz ism-familiyalarini nusxalab yozing va botga yuboring. Bot ularga login-parol yaratib qaytaradi.\n\n"
        "<b>3. Test jarayonini o'tkazish:</b>\n"
        "Testni sozlab bo'lgach, kodni o'quvchilarga tarqating va darsda hamma tayyor bo'lganda Start 🚀 tugmasini bosing.\n\n"
        "<b>4. Bepul Demo va Taklif tizimi:</b>\n"
        "Barcha yangi o'qituvchilarga 1 ta bepul test yaratish limiti beriladi. Limit tugagach, botga yangi 5 ta hamkasb ustozni taklif qilib, <b>yana 1 ta bepul test limiti</b> olishingiz mumkin! Buning uchun 'Mening obunam' bo'limidan taklif havolangizni oling.\n\n"
        "Batafsil rasmli qo'llanma va shablonlarni yuklab olish:"
    )
    await message.answer(guide_text, reply_markup=get_guide_download_kb(), parse_mode="HTML")


@router.callback_query(F.data == "download_template_docx")
async def cb_download_docx(callback: CallbackQuery):
    """Namunali Word shablonini dinamik hosil qilib jo'natish"""
    import docx

    doc = docx.Document()
    doc.add_heading("Test Savollari Shablon (Namunaviy)", 0)
    doc.add_paragraph(
        "Diqqat: Ushbu jadvaldagi ustunlar soni (aynan 6 ta) va tartibi o'zgarmasligi shart! "
        "Formulalar yoki chizmalarni 'Savol' katagiga rasm qilib joylang."
    )

    # 3 qator va 6 ustunli jadval
    table = doc.add_table(rows=3, cols=6)
    table.style = "Table Grid"

    # Sarlavhalar
    headers = ["T/r", "Savol", "To'g'ri javob", "Noto'g'ri javob 1", "Noto'g'ri javob 2", "Noto'g'ri javob 3"]
    for i, name in enumerate(headers):
        table.rows[0].cells[i].text = name

    # Namunalar
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "O'zbekiston Respublikasining poytaxti qaysi shahar?"
    table.rows[1].cells[2].text = "Toshkent"
    table.rows[1].cells[3].text = "Samarqand"
    table.rows[1].cells[4].text = "Buxoro"
    table.rows[1].cells[5].text = "Andijon"

    table.rows[2].cells[0].text = "2"
    table.rows[2].cells[1].text = "Matematik masalani yeching: 25 * 4 = ?"
    table.rows[2].cells[2].text = "100"
    table.rows[2].cells[3].text = "90"
    table.rows[2].cells[4].text = "80"
    table.rows[2].cells[5].text = "120"

    stream = io.BytesIO()
    doc.save(stream)
    file_bytes = stream.getvalue()

    await callback.message.answer_document(
        BufferedInputFile(file_bytes, filename="savollar_shabloni.docx"),
        caption="📝 Word savollar shabloni. Ushbu namunani to'ldirib yuboring.",
    )
    await callback.answer()


@router.callback_query(F.data == "download_template_xlsx")
async def cb_download_xlsx(callback: CallbackQuery):
    """Namunali Excel student shablonini dinamik hosil qilib jo'natish"""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Studentlar"
    ws.views.sheetView[0].showGridLines = True

    # Sarlavha
    ws["A1"] = "Ism Familiya"
    ws["A1"].font = Font(name="Calibri", size=11, bold=True)

    # Namunalar
    ws["A2"] = "Ali Valiyev"
    ws["A3"] = "Kamola Umarova"
    ws["A4"] = "Eshmat Toshmatov"

    stream = io.BytesIO()
    wb.save(stream)
    file_bytes = stream.getvalue()

    await callback.message.answer_document(
        BufferedInputFile(file_bytes, filename="studentlar_shabloni.xlsx"),
        caption="📊 Studentlar ro'yxati shabloni. Faqat A ustunini to'ldiring.",
    )
    await callback.answer()


@router.callback_query(F.data == "download_guide_pdf")
async def cb_download_guide_pdf(callback: CallbackQuery):
    """Batafsil foydalanish yo'riqnomasi (.docx formatda, foydalanuvchiga oson bo'lishi uchun)"""
    import docx

    doc = docx.Document()
    doc.add_heading("Moderatorlar uchun Yo'riqnoma", 0)

    doc.add_heading("1. Word faylini to'ldirish", 1)
    doc.add_paragraph(
        "Yuklab olingan 'savollar_shabloni.docx' faylini oching. Jadval 6 ta ustundan iborat bo'lishi kerak. "
        "Har bir qatorga savol yozib, variantlarni to'ldiring. Jadval tashqarisidagi har qanday sarlavhalar va logotiplar bot tomonidan hisobga olinmaydi."
    )

    doc.add_heading("2. Excel o'quvchilar ro'yxati", 1)
    doc.add_paragraph(
        "O'quvchilarning ro'yxatini A ustuniga (Ism Familiya) nusxalab yozing. Bot sizga ularning login va parollarini yaratib beradi. "
        "Bitta telefondan faqat bir kishi faol test yechishi mumkin. O'quvchi testni yechib bo'lishi bilanoq tizimdan avtomat chiqib ketadi (Logout) "
        "va telefon keyingi o'quvchiga bo'shatiladi."
    )

    doc.add_heading("3. Imtihonni boshlash va nazorat qilish", 1)
    doc.add_paragraph(
        "Kodni o'quvchilarga tarqating. Ular botga kod va o'z login-paroli bilan kirib kutish zalida (Lobby) turishadi. "
        "Hamma kirganidan so'ng 'Start' bosing. Imtihon davomida 'Jonli Nazorat' bo'limidan kim nechanchi savolda turganini onlayn kuzatib boring."
    )

    stream = io.BytesIO()
    doc.save(stream)
    file_bytes = stream.getvalue()

    await callback.message.answer_document(
        BufferedInputFile(file_bytes, filename="moderator_manual.docx"),
        caption="📄 Moderatorlik qo'llanmasi (Word hujjat).",
    )
    await callback.answer()


# ==========================================
# 4. CHIQISH (LOGOUT)
# ==========================================
@router.message(F.text == "Tizimdan chiqish")
async def process_logout(message: Message, state: FSMContext):
    """Moderator hisobidan chiqish. Agar 0 ta testi bo'lsa profilini o'chiradi (mehmonga aylanadi)"""
    await state.clear()
    user_id = message.from_user.id

    async with async_session() as session:
        # Testlari sonini tekshirish
        quiz_count_stmt = select(func.count(Quiz.id)).where(Quiz.teacher_id == user_id)
        quiz_res = await session.execute(quiz_count_stmt)
        quiz_count = quiz_res.scalar()

        if quiz_count == 0 and user_id not in ADMIN_IDS:
            # Testi yo'q va admin emas - bazadan butunlay o'chiramiz (curious student)
            del_stmt = delete(User).where(User.id == user_id)
            await session.execute(del_stmt)
            await session.commit()
            await message.answer(
                "Moderator hisobingiz tozalab o'chirildi. Bosh sahifaga qaytdingiz.",
                reply_markup=get_welcome_kb(),
            )
        else:
            # Testi bor yoki Bosh Admin - ma'lumotlarni o'chirmaymiz.
            if user_id in ADMIN_IDS:
                await message.answer(
                    "Tizimdan chiqdingiz. Admin paneliga qaytdingiz.",
                    reply_markup=get_admin_menu(),
                )
            else:
                await message.answer(
                    "Tizimdan chiqdingiz. Testlaringiz va ma'lumotlaringiz saqlanib qoldi.\n"
                    "Qayta kirish uchun istalgan payt 'Moderator bo'lish' tugmasini bossangiz bas.",
                    reply_markup=get_welcome_kb(),
                )


# ==========================================
# 5. STUDENTGA QAYTISH
# ==========================================
@router.message(F.text == "Talaba rejimiga o'tish")
async def process_teacher_to_student(message: Message, state: FSMContext):
    """Moderator kabinetidan student login oynasiga o'tish (FSM state ni o'zgartirish)"""
    await state.clear()
    await state.set_state(AuthStates.waiting_for_student_login)
    await message.answer(
        "Student sifatida kirish uchun Moderatoringiz bergan **Login**ni kiriting:",
        reply_markup=get_student_menu(),
        parse_mode="Markdown",
    )


# ==========================================
# 6. ADMINGA MUROJAAT (FEEDBACK)
# ==========================================
@router.message(F.text == "Adminga murojaat")
async def teacher_support_start(message: Message, state: FSMContext):
    """Adminga murojaat yozish holatiga o'tkazish"""
    await state.set_state(TeacherStates.waiting_for_support_message)
    await message.answer(
        "✉️ **Bosh adminga murojaat yo'llang:**\n\n"
        "Takliflaringiz yoki texnik muammolarni batafsil yozib yuboring:",
        parse_mode="Markdown",
    )


@router.callback_query(F.data == "contact_admin")
async def cb_contact_admin(callback: CallbackQuery, state: FSMContext):
    await state.set_state(TeacherStates.waiting_for_support_message)
    await callback.message.answer(
        "✉️ **Bosh adminga murojaat yo'llang:**\n\n"
        "Takliflaringiz yoki texnik muammolarni batafsil yozib yuboring:",
        parse_mode="Markdown",
    )
    await callback.answer()


@router.message(TeacherStates.waiting_for_support_message)
async def process_support_msg(message: Message, state: FSMContext, bot: Bot):
    """Yozilgan murojaatni Super Adminlarga inline javob yozish tugmasi bilan yuborish"""
    support_text = message.text.strip()
    if not support_text:
        await message.answer("Murojaat matni bo'sh bo'lishi mumkin emas.")
        return

    user_id = message.from_user.id
    full_name = message.from_user.full_name
    username = message.from_user.username

    # Super Adminlarga jo'natish
    from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

    admin_kb = InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="✍️ Javob yozish", callback_data=f"reply_ticket_{user_id}")]
        ]
    )

    admin_msg = (
        f"💬 **YANGI MUROJAAT KELDI**\n\n"
        f"👤 **Yuboruvchi:** {full_name}\n"
        f"🆔 **ID:** `{user_id}`\n"
        f"🔗 **Username:** @{username if username else 'yoq'}\n"
        f"💬 **Murojaat matni:**\n{support_text}"
    )

    for admin_id in ADMIN_IDS:
        try:
            await bot.send_message(admin_id, admin_msg, reply_markup=admin_kb, parse_mode="Markdown")
        except Exception:
            pass

    await state.clear()
    await message.answer(
        "✅ Murojaatingiz bosh adminga yetkazildi. Tez orada javob olasiz!",
        reply_markup=get_teacher_menu(is_demo=False),
    )


# ==========================================
# 7. MENING OBUNAM (SUBSCRIPTION)
# ==========================================
@router.message(F.text == "Mening obunam")
async def show_subscription(message: Message):
    """Moderatorning joriy obunasi va tarif limitlari haqida ma'lumot berish"""
    user_id = message.from_user.id

    async with async_session() as session:
        # Joriy obunalarni olish
        sub_stmt = select(Subscription).where(Subscription.user_id == user_id)
        sub_res = await session.execute(sub_stmt)
        subs = sub_res.scalars().all()

        # Tizim narxlari va limitlarini olish
        settings_res = await session.execute(select(SystemSettings))
        settings = settings_res.scalar()

        # Obunalarni guruhlash
        has_premium = any(s.type == "monthly" and s.expires_at > local_now() for s in subs)
        premium_sub = next((s for s in subs if s.type == "monthly" and s.expires_at > local_now()), None)

        onetime_credits = sum(s.credits for s in subs if s.type == "onetime")

        # Matnni shakllantirish
        status_text = "🎟 **SIZNING FAOLLIK TARIFLARINGIZ:**\n\n"

        if has_premium:
            status_text += (
                f"🌟 **Premium Obuna (1 oylik):** FAOL ✅\n"
                f"📅 Tugash muddati: {premium_sub.expires_at.strftime('%Y-%m-%d %H:%M')}\n"
                f"👥 Ruxsat etilgan studentlar soni: *{settings.monthly_max_students} ta* (har bir testda)\n"
                f"📁 Test yaratish limiti: *Cheksiz*\n\n"
            )
        else:
            status_text += f"🌟 **Premium Obuna:** Faol emas ❌\n\n"

        if onetime_credits > 0:
            status_text += (
                f"🎫 **Bir martalik limitlar:** FAOL ✅\n"
                f"💳 Qolgan limitlar soni: *{onetime_credits} ta test uchun*\n"
                f"👥 Ruxsat etilgan studentlar soni: *{settings.onetime_max_students} ta* (har bir testda)\n\n"
            )
        else:
            status_text += f"🎫 **Bir martalik limitlar:** Mavjud emas ❌\n\n"

        if not has_premium and onetime_credits <= 0:
            bot_info = await message.bot.get_me()
            ref_link = f"https://t.me/{bot_info.username}?start=ref_{user_id}"
            user_stmt = select(User).where(User.id == user_id)
            db_u = (await session.execute(user_stmt)).scalar_one_or_none()
            ref_count = db_u.referral_count if db_u else 0

            status_text += (
                f"🎁 **Bepul Demo Paket:** FAOL ✅\n"
                f"👥 Ruxsat etilgan studentlar soni: *{settings.demo_max_students} ta* (har bir testda)\n"
                f"📁 Test yaratish limiti: *1 ta test yaratish imkoniyati*\n"
                f"🤝 **Bepul limit olish:**\n"
                f"Quyidagi taklif havolasini 5 ta hamkasb ustozga yuboring. Ular botimizga kirib `/start` tugmasini bosishi bilanoq sizga **1 ta bepul test yaratish limiti** taqdim etiladi.\n"
                f"🔗 Taklif havolangiz: `{ref_link}`\n"
                f"📊 Taklif ko'rsatkichi: *{ref_count}/5*\n"
                f"_Ko'proq test o'tkazish uchun Premium yoki Limit xarid qiling!_\n\n"
            )

        # Dynamic tariff pricing info
        status_text += (
            f"📌 **Tariflar Narxlari (Xarid qilish uchun Adminga murojaat qiling):**\n"
            f"- 🎫 1 martalik test limiti: *{settings.onetime_price:,} so'm* (Maks. {settings.onetime_max_students} talaba)\n"
            f"- 🌟 1 oylik Premium obuna: *{settings.monthly_price:,} so'm* (Maks. {settings.monthly_max_students} talaba, cheksiz testlar)\n"
        )

    await message.answer(status_text, reply_markup=get_teacher_menu(is_demo=False), parse_mode="Markdown")


# ==========================================
# 8. YANGI TEST YARATISH (WIZARD)
# ==========================================
@router.message(F.text == "Yangi test yaratish")
async def create_quiz_start(message: Message, state: FSMContext):
    """Test yaratishni boshlash (demo limitlarini tekshirish)"""
    user_id = message.from_user.id

    async with async_session() as session:
        # Tariflarni tekshirish (credits va active subscriptions)
        sub_stmt = select(Subscription).where(Subscription.user_id == user_id)
        sub_res = await session.execute(sub_stmt)
        subs = sub_res.scalars().all()

        is_premium = any(s.type == "monthly" and s.expires_at > local_now() for s in subs)
        has_credits = any(s.type == "onetime" and s.credits > 0 for s in subs)

        # Hozirgi testlar sonini olish
        quiz_count_stmt = select(func.count(Quiz.id)).where(Quiz.teacher_id == user_id)
        quiz_count_res = await session.execute(quiz_count_stmt)
        quiz_count = quiz_count_res.scalar()

        if user_id not in ADMIN_IDS and not is_premium and not has_credits:
            # Agar faqat bepul demo bo'lsa va 1 ta test yaratilgan bo'lsa
            if quiz_count >= 1:
                bot_info = await message.bot.get_me()
                ref_link = f"https://t.me/{bot_info.username}?start=ref_{user_id}"
                
                # Taklif qilingan foydalanuvchilar sonini bazadan olamiz
                user_stmt = select(User).where(User.id == user_id)
                db_u = (await session.execute(user_stmt)).scalar_one_or_none()
                ref_count = db_u.referral_count if db_u else 0
                
                await message.answer(
                    f"⚠️ **Hurmatli ustoz, botimiz pullik va professional tizim hisoblanadi.**\n\n"
                    f"Siz botni bepul \"Demo\" darajasida ishlatib bo'ldingiz. Yangi test yaratish uchun "
                    f"Premium obuna yoki 1 martalik limit xarid qilishingiz lozim.\n\n"
                    f"🎁 **Bepul muqobil yo'l:**\n"
                    f"Quyidagi taklif havolasini **5 ta hamkasbingizga** yuboring. Ular botimizga kirib "
                    f"`/start` tugmasini bosgan zahoti sizga **1 ta bepul limit** taqdim etiladi.\n\n"
                    f"🔗 Sizning taklif havolangiz:\n{ref_link}\n\n"
                    f"📊 Taklif qilingan yangi hamkasblar: *{ref_count}/5*",
                    parse_mode="Markdown",
                    disable_web_page_preview=True
                )
                return

    await state.set_state(TeacherStates.waiting_for_quiz_title)
    await message.answer("📝 **Yangi test yaratish:**\n\nTest nomini kiriting (Masalan: Matematika 5-sinf):")


@router.message(TeacherStates.waiting_for_quiz_title)
async def process_quiz_title(message: Message, state: FSMContext):
    title = message.text.strip()
    if not title or len(title) < 3:
        await message.answer("Iltimos, yaroqliroq test nomini kiriting:")
        return

    await state.update_data(quiz_title=title)
    await state.set_state(TeacherStates.waiting_for_quiz_description)
    await message.answer("📝 Test tavsifini yozing (Masalan: 1-smena uchun nazorat ishi):")


@router.message(TeacherStates.waiting_for_quiz_description)
async def process_quiz_description(message: Message, state: FSMContext):
    desc = message.text.strip()
    await state.update_data(quiz_description=desc, uploaded_variants=[])

    # Word yuklash holatiga o'tish (Imtihon davomiyligi uni boshlashdan oldin so'raladi)
    await state.set_state(TeacherStates.waiting_for_docx_file)
    await message.answer(
        "📝 **1-variant uchun Word (.docx) test savollarini yuklang:**\n\n"
        "_Fayl sarlavha ostidagi jadval 6 ta ustundan iborat ekanligiga ishonch hosil qiling._",
        parse_mode="Markdown",
    )


@router.message(TeacherStates.waiting_for_docx_file, F.document)
async def process_quiz_docx(message: Message, state: FSMContext, bot: Bot):
    """Yuklangan docx test faylini parse va validatsiya qilish"""
    doc = message.document
    if not doc.file_name.endswith(".docx"):
        await message.answer("Iltimos, faqat Word (.docx) kengaytmali fayl yuklang.")
        return

    # Faylni yuklab olish
    file_bytes = io.BytesIO()
    await bot.download(doc.file_id, destination=file_bytes)
    file_bytes.seek(0)

    # Word jadvalini parse qilish
    result = parse_docx_questions(file_bytes)
    if "error" in result:
        await message.answer(f"{result['error']}\n\nIltimos, xatoni to'g'rilab, faylni qaytadan yuboring:")
        return

    questions = result["questions"]
    data = await state.get_data()
    uploaded_variants = data.get("uploaded_variants", [])

    # Joriy variant tartibini aniqlash (1, 2, 3...)
    current_variant_idx = len(uploaded_variants) + 1
    variant_name = f"Variant {current_variant_idx}"

    # Variant ma'lumotlarini holatda (FSM) saqlash
    uploaded_variants.append({"variant": str(current_variant_idx), "questions": questions})
    await state.update_data(uploaded_variants=uploaded_variants)

    await message.answer(
        f"✅ **{variant_name} muvaffaqiyatli o'qildi!**\n"
        f"Jami aniqlangan savollar soni: {len(questions)} ta.\n\n"
        f"❓ **Yana qo'shimcha variant yuklaysizmi?**\n"
        f"Agar yuklamoqchi bo'lsangiz, keyingi variant uchun Word faylini yuboring.\n"
        f"Agar variantlar tugagan bo'lsa, pastdagi tugmani bosing:",
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[
                [
                    InlineKeyboardButton(
                        text="Tugatish va O'quvchilarni yuklash ➡️",
                        callback_data="docx_finished",
                    )
                ]
            ]
        ),
        parse_mode="Markdown",
    )


@router.callback_query(TeacherStates.waiting_for_docx_file, F.data == "docx_finished")
async def cb_docx_finished(callback: CallbackQuery, state: FSMContext):
    """Word yuklash yakunlandi, endi o'quvchilar ro'yxatini yuklashni so'raymiz"""
    data = await state.get_data()
    uploaded_variants = data.get("uploaded_variants", [])

    if not uploaded_variants:
        await callback.answer("Hech bo'lmasa 1 ta variant yuklanishi shart!", show_alert=True)
        return

    await state.set_state(TeacherStates.waiting_for_excel_file)
    await callback.message.edit_text(
        "📊 **Endi studentlar ism-familiyalari ro'yxatini (Excel - .xlsx) yuklang:**\n\n"
        "_Faylning birinchi ustun sarlavhasi 'Ism Familiya' bo'lishi shart._",
        parse_mode="Markdown",
    )
    await callback.answer()


@router.message(TeacherStates.waiting_for_excel_file, F.document)
async def process_quiz_excel(message: Message, state: FSMContext, bot: Bot):
    """Yuklangan Excel o'quvchilar ro'yxatini tekshirish va ma'lumotlar bazasiga testni saqlash (Tranzaksiya bilan)"""
    doc = message.document
    if not doc.file_name.endswith(".xlsx"):
        await message.answer("Iltimos, faqat Excel (.xlsx) kengaytmali fayl yuklang.")
        return

    # Faylni yuklab olish
    file_bytes = io.BytesIO()
    await bot.download(doc.file_id, destination=file_bytes)
    file_bytes.seek(0)

    # Excel o'quvchilarni parse qilish
    result = parse_students_excel(file_bytes)
    if "error" in result:
        await message.answer(f"{result['error']}\n\nIltimos, xatoni to'g'rilab, faylni qaytadan yuboring:")
        return

    students_credentials = result["students"]
    user_id = message.from_user.id

    data = await state.get_data()
    title = data["quiz_title"]
    desc = data["quiz_description"]
    uploaded_variants = data["uploaded_variants"]

    # Studentlar limiti tekshiruvi (Tarif bo'yicha)
    async with async_session() as session:
        # Tarif ma'lumotlarini olish
        sub_stmt = select(Subscription).where(Subscription.user_id == user_id)
        sub_res = await session.execute(sub_stmt)
        subs = sub_res.scalars().all()

        is_premium = any(s.type == "monthly" and s.expires_at > local_now() for s in subs)
        has_credits = any(s.type == "onetime" and s.credits > 0 for s in subs)
        is_demo_active = any(s.type == "free_demo" and s.expires_at > local_now() for s in subs)

        # Cheklovlarni bazadan olish
        settings_res = await session.execute(select(SystemSettings))
        settings = settings_res.scalar()

        if is_premium:
            selected_sub = next((s for s in subs if s.type == "monthly" and s.expires_at > local_now()), None)
            max_students = selected_sub.max_students_limit if selected_sub and selected_sub.max_students_limit else settings.monthly_max_students
            sub_type = "monthly"
        elif has_credits:
            selected_sub = next((s for s in subs if s.type == "onetime" and s.credits > 0), None)
            max_students = selected_sub.max_students_limit if selected_sub and selected_sub.max_students_limit else settings.onetime_max_students
            sub_type = "onetime"
        elif is_demo_active:
            selected_sub = next((s for s in subs if s.type == "free_demo" and s.expires_at > local_now()), None)
            max_students = selected_sub.max_students_limit if selected_sub and selected_sub.max_students_limit else settings.demo_max_students
            sub_type = "free_demo"
        else:
            await message.answer(
                "❌ **Sizning barcha obunalaringiz va demo muddatingiz tugagan!**\n"
                "Test yaratish uchun yangi obuna yoki limit xarid qiling.",
                parse_mode="Markdown"
            )
            return

        if len(students_credentials) > max_students:
            await message.answer(
                f"⚠️ **Limit cheklovi!**\nSizning tarigingizda maksimal student soni: *{max_students} ta*.\n"
                f"Siz yuklagan ro'yxatda esa: *{len(students_credentials)} ta* o'quvchi bor.\n\n"
                f"Iltimos, o'quvchilar sonini kamaytirib yuklang yoki obunangizni yangilang.",
                parse_mode="Markdown",
            )
            return

        # Ma'lumotlar bazasiga tranzaksiya bilan saqlash
        # Test kodini generatsiya qilish
        while True:
            code = "".join(random.choices(string.digits, k=6))
            code_check = await session.execute(select(Quiz).where(Quiz.code == code))
            if not code_check.scalar_one_or_none():
                break

        # Faol variantlar ro'yxati
        active_vars = [v["variant"] for v in uploaded_variants]

        new_quiz = Quiz(
            teacher_id=user_id,
            title=title,
            description=desc,
            duration_minutes=30,  # Boshlash arafasida o'zgartiriladigan default qiymat
            code=code,
            active_variants=active_vars,
            status="waiting",
        )
        session.add(new_quiz)
        await session.flush()  # ID sini olish uchun flush

        # Savollarni saqlash
        for var_data in uploaded_variants:
            var_name = var_data["variant"]
            for q in var_data["questions"]:
                new_q = Question(
                    quiz_id=new_quiz.id,
                    question_text=q["question_text"],
                    image_data=q["image_data"],
                    correct_answer=q["correct_answer"],
                    distractors=q["distractors"],
                    variant=var_name,
                )
                session.add(new_q)

        # Studentlarni saqlash (parollarni bcrypt bilan hash qilamiz)
        import bcrypt
        for s in students_credentials:
            salt = bcrypt.gensalt()
            hashed_pw = bcrypt.hashpw(s["password"].encode(), salt).decode("utf-8")
            new_s = Student(
                quiz_id=new_quiz.id,
                full_name=s["full_name"],
                login=s["login"],
                password=hashed_pw,
            )
            session.add(new_s)

        # Limitdan 1 ta ayirish (agar bir martalik bo'lsa)
        if sub_type == "onetime" and selected_sub:
            selected_sub.credits -= 1

        # Referal hisoblagichni nolga tushirish (agar barcha limitlar tugagan bo'lsa)
        quiz_count_stmt = select(func.count(Quiz.id)).where(Quiz.teacher_id == user_id)
        quiz_count_res = await session.execute(quiz_count_stmt)
        quiz_count = quiz_count_res.scalar() or 0

        # Obunalarni qayta tekshiramiz
        sub_stmt2 = select(Subscription).where(Subscription.user_id == user_id)
        sub_res2 = await session.execute(sub_stmt2)
        subs2 = sub_res2.scalars().all()
        
        is_prem = any(s.type == "monthly" and s.expires_at > local_now() for s in subs2)
        onetime_credits_left = sum(s.credits for s in subs2 if s.type == "onetime")
        
        # Agar u premium bo'lmasa va 0 ta onetime limiti qolgan bo'lsa (shuningdek demo test ishlatilgan bo'lsa)
        if not is_prem and onetime_credits_left <= 0:
            user_stmt = select(User).where(User.id == user_id)
            user_res = await session.execute(user_stmt)
            user = user_res.scalar_one_or_none()
            if user:
                user.referral_count = 0

        await session.commit()

    # O'qituvchi uchun login-parollar Excelini generatsiya qilish va yuborish
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Student Logins"
    ws.views.sheetView[0].showGridLines = True

    ws["A1"] = "Ism Familiya"
    ws["B1"] = "Login"
    ws["C1"] = "Parol"
    for col_idx in range(1, 4):
        ws.cell(row=1, column=col_idx).font = Font(name="Calibri", size=11, bold=True)

    for idx, s in enumerate(students_credentials, 2):
        ws.cell(row=idx, column=1, value=s["full_name"])
        ws.cell(row=idx, column=2, value=s["login"])
        ws.cell(row=idx, column=3, value=s["password"])

    ws.column_dimensions["A"].width = 25
    ws.column_dimensions["B"].width = 18
    ws.column_dimensions["C"].width = 12

    stream = io.BytesIO()
    wb.save(stream)
    excel_bytes = stream.getvalue()

    await state.clear()
    await message.answer(
        f"🚀 **Test muvaffaqiyatli yaratildi!**\n\n"
        f"🔑 **Test kodi:** `{code}`\n"
        f"⏳ **Vaqti:** {new_quiz.duration_minutes} daqiqa\n"
        f"📊 **Studentlar soni:** {len(students_credentials)} ta\n"
        f"📁 **Yuklangan variantlar:** {', '.join(active_vars)}\n\n"
        f"_Quyida studentlarning login va parollari yozilgan Excel varaq yuborildi. "
        f"Buni o'quvchilarga tarqating yoki chop eting._",
        reply_markup=get_teacher_menu(is_demo=False),
        parse_mode="Markdown",
    )

    await message.answer_document(
        BufferedInputFile(excel_bytes, filename=f"logins_{code}.xlsx"),
        caption=f"🔑 Test kodi: {code} uchun student login-parollari.",
    )


# ==========================================
# 9. MENING TESTLARIM (BOSHQARUV)
# ==========================================
@router.message(F.text == "Mening testlarim")
async def show_my_quizzes(message: Message, state: FSMContext):
    """Moderatorning barcha yaratgan testlari ro'yxatini chiqaradi"""
    user_id = message.from_user.id

    async with async_session() as session:
        stmt = (
            select(Quiz)
            .where(Quiz.teacher_id == user_id)
            .order_by(Quiz.created_at.desc())
        )
        res = await session.execute(stmt)
        quizzes = res.scalars().all()

        if not quizzes:
            await message.answer(
                "Sizda hali yaratilgan testlar mavjud emas. Yangi test yaratish uchun pastdagi tugmani bosing."
            )
            return

        response_text = "📚 **Sizning testlaringiz ro'yxati:**\n\n"
        buttons = []

        from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

        for q in quizzes:
            status_emoji = "⏳" if q.status == "waiting" else ("🚀" if q.status == "active" else "✅")
            status_text = "Kutilmoqda" if q.status == "waiting" else ("Faol" if q.status == "active" else "Yakunlandi")

            response_text += (
                f"{status_emoji} *{q.title}*\n"
                f"└ Kodi: `{q.code}` | Holati: {status_text}\n\n"
            )

            # Testni boshqarish uchun inline tugma
            buttons.append(
                [
                    InlineKeyboardButton(
                        text=f"Boshqarish: {q.title[:20]}...",
                        callback_data=f"manage_quiz_{q.id}",
                    )
                ]
            )

        await message.answer(
            response_text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
            parse_mode="Markdown",
        )


@router.callback_query(F.data.startswith("manage_quiz_"))
async def cb_manage_quiz(callback: CallbackQuery, state: FSMContext):
    """Testni boshqarish panelini ochadi (variantlar, boshlash, kutish zali, natijalar)"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz:
            await callback.answer("Ushbu test topilmadi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!", show_alert=True)
            return

        # Studentlar sonini hisoblash
        stud_count_stmt = select(func.count(Student.id)).where(Student.quiz_id == quiz_id)
        stud_count_res = await session.execute(stud_count_stmt)
        stud_count = stud_count_res.scalar()

        # Lobby kirgan studentlar
        logged_stmt = select(func.count(Student.id)).where(
            Student.quiz_id == quiz_id, Student.telegram_id.isnot(None)
        )
        logged_res = await session.execute(logged_stmt)
        logged_count = logged_res.scalar()

        status_text = "Kutilmoqda" if quiz.status == "waiting" else ("Faol" if quiz.status == "active" else "Yakunlandi")

        control_text = (
            f"⚙️ **TESTNI BOSHQARISH PANELI**\n\n"
            f"📝 **Nomi:** {quiz.title}\n"
            f"🔑 **Kirish kodi:** `{quiz.code}`\n"
            f"⏳ **Davomiyligi:** {quiz.duration_minutes} daqiqa\n"
            f"📂 **Variantlar:** {', '.join(quiz.active_variants)}\n"
            f"👥 **Studentlar:** {logged_count}/{stud_count} ta kirdi\n"
            f"📊 **Holati:** {status_text}\n"
        )

        from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

        buttons = []

        if quiz.status == "waiting":
            # Imtihon boshlanmagan holat
            buttons.append(
                [
                    InlineKeyboardButton(
                        text="👥 Kutish Zali (Lobby)",
                        callback_data=f"lobby_monitor_{quiz_id}",
                    )
                ]
            )
            buttons.append(
                [
                    InlineKeyboardButton(
                        text="🚀 Imtihonni boshlash",
                        callback_data=f"launch_quiz_{quiz_id}",
                    )
                ]
            )
        elif quiz.status == "active":
            # Imtihon faol holat
            buttons.append(
                [
                    InlineKeyboardButton(
                        text="📈 Jonli Nazorat (Monitor)",
                        callback_data=f"live_monitor_{quiz_id}",
                    )
                ]
            )
            buttons.append(
                [
                    InlineKeyboardButton(
                        text="🛑 Imtihonni yakunlash",
                        callback_data=f"terminate_quiz_{quiz_id}",
                    )
                ]
            )
        elif quiz.status == "completed":
            # Imtihon yakunlangan holat
            buttons.append(
                [
                    InlineKeyboardButton(
                        text="📊 Natijalarni yuklash (Excel)",
                        callback_data=f"download_results_{quiz_id}",
                    )
                ]
            )

        # O'chirish tugmasi
        buttons.append(
            [
                InlineKeyboardButton(
                    text="❌ Testni butunlay o'chirish",
                    callback_data=f"delete_quiz_{quiz_id}",
                )
            ]
        )

        await callback.message.edit_text(
            control_text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
            parse_mode="Markdown",
        )
        await callback.answer()


# ==========================================
# 7. KUTISH ZALI VA IMTIHONNI BOSHLASH
# ==========================================
@router.callback_query(F.data.startswith("lobby_monitor_"))
async def cb_lobby_monitor(callback: CallbackQuery):
    """Kutish zalini (tayyor va kirmagan o'quvchilar ro'yxati) ko'rsatish"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz:
            await callback.answer("Test topilmadi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!", show_alert=True)
            return

        # Barcha studentlarni olish
        students = await session.scalars(
            select(Student).where(Student.quiz_id == quiz_id).order_by(Student.full_name)
        )
        students = students.all()

        entered_students = [s for s in students if s.telegram_id is not None]
        missing_students = [s for s in students if s.telegram_id is None]

        lobby_text = (
            f"👥 **KUTISH ZALI: {quiz.title}**\n"
            f"🔑 Kodi: `{quiz.code}`\n"
            f"Statistika: {len(entered_students)}/{len(students)} ta o'quvchi kirdi.\n\n"
            f"✅ **KIRGANLAR ({len(entered_students)}):**\n"
        )
        for s in entered_students:
            lobby_text += f"- {s.full_name}\n"

        lobby_text += f"\n❌ **KIRMAGANLAR ({len(missing_students)}):**\n"
        for s in missing_students:
            lobby_text += f"- {s.full_name} (login: `{s.login}`)\n"

        from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

        buttons = [
            [
                InlineKeyboardButton(
                    text="🔄 Yangilash", callback_data=f"lobby_monitor_{quiz_id}"
                ),
                InlineKeyboardButton(
                    text="🚀 Boshlash", callback_data=f"launch_quiz_{quiz_id}"
                ),
            ],
            [InlineKeyboardButton(text="🔙 Orqaga", callback_data=f"manage_quiz_{quiz_id}")],
        ]

        await callback.message.edit_text(
            lobby_text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
            parse_mode="Markdown",
        )
        await callback.answer()


@router.callback_query(F.data.startswith("launch_quiz_"))
async def cb_launch_quiz(callback: CallbackQuery, state: FSMContext):
    """Imtihonni boshlashdan oldin davomiylikni daqiqalarda so'rash"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz or quiz.status != "waiting":
            await callback.answer("Imtihonni boshlab bo'lmaydi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!", show_alert=True)
            return

    await state.update_data(launch_quiz_id=quiz_id)
    await state.set_state(TeacherStates.waiting_for_start_duration)

    await callback.message.edit_text(
        "⏳ **Imtihon davomiyligini kiriting (daqiqalarda, masalan: 45, 60 kabi faqat musbat butun son yozing):**\n\n"
        "_Kiritilgan vaqt tugashi bilan o'quvchilar uchun test avtomatik yakunlanadi._",
        reply_markup=InlineKeyboardMarkup(
            inline_keyboard=[
                [InlineKeyboardButton(text="🔙 Bekor qilish", callback_data=f"manage_quiz_{quiz_id}")]
            ]
        ),
        parse_mode="Markdown",
    )
    await callback.answer()


@router.message(TeacherStates.waiting_for_start_duration)
async def process_start_duration(message: Message, state: FSMContext, bot: Bot):
    """Moderator kiritgan davomiylikni olib testni ishga tushirish (Start)"""
    duration_raw = message.text.strip()
    if not duration_raw.isdigit() or int(duration_raw) <= 0:
        await message.answer("Iltimos, faqat musbat butun son yozing (Masalan: 45):")
        return

    duration = int(duration_raw)
    state_data = await state.get_data()
    quiz_id = state_data.get("launch_quiz_id")

    if not quiz_id:
        await message.answer("Xatolik: Imtihon ID si topilmadi. Iltimos, qaytadan urinib ko'ring.")
        await state.clear()
        return

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz or quiz.status != "waiting":
            await message.answer("Imtihonni boshlab bo'lmaydi!")
            await state.clear()
            return

        if quiz.teacher_id != message.from_user.id and message.from_user.id not in ADMIN_IDS:
            await message.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!")
            await state.clear()
            return

        # Imtihonni faollashtirish va davomiyligini belgilash
        quiz.status = "active"
        quiz.duration_minutes = duration
        quiz.start_time = local_now()
        quiz.end_time = quiz.start_time + datetime.timedelta(minutes=duration)

        # Kirgan studentlarni olish (FSM dagi waiting o'quvchilarga xabar berish)
        students = await session.scalars(
            select(Student).where(Student.quiz_id == quiz_id, Student.telegram_id.isnot(None))
        )
        students = students.all()

        await session.commit()

    # FSM holatini tozalash
    await state.clear()

    # Barcha studentlarga test boshlangani haqida signal jo'natish
    from keyboards.bot_keyboards import get_student_start_exam_kb

    success_notified = 0
    for s in students:
        try:
            await bot.send_message(
                s.telegram_id,
                f"🚀 **IMTIHON BOSHLANDI!**\n\n"
                f"📚 Test: {quiz.title}\n"
                f"⏳ Davomiyligi: {duration} daqiqa\n\n"
                f"Testni boshlash uchun quyidagi tugmani bosing:",
                reply_markup=get_student_start_exam_kb(quiz.code),
                parse_mode="Markdown",
            )
            success_notified += 1
        except Exception:
            pass

    await message.answer(
        f"🚀 **Imtihon faollashtirildi!**\n"
        f"Imtihon davomiyligi: *{duration} daqiqa*.\n"
        f"Lobbyda kutib turgan {success_notified} ta o'quvchiga start xabarlari yuborildi."
    )

    # Jonli nazoratni ochish uchun callback soxtalashtirish va chaqirish
    from aiogram.types import Chat, User as TelegramUser
    dummy_chat = Chat(id=message.chat.id, type="private")
    dummy_user = TelegramUser(id=message.from_user.id, is_bot=False, first_name=message.from_user.first_name)
    dummy_msg = Message(
        message_id=message.message_id,
        date=message.date,
        chat=dummy_chat,
        from_user=dummy_user,
        text="",
    )
    dummy_callback = CallbackQuery(
        id="dummy",
        from_user=dummy_user,
        chat_instance="dummy",
        message=dummy_msg,
        data=f"live_monitor_{quiz_id}",
    )
    await cb_live_monitor(dummy_callback)


# ==========================================
# 8. JONLI IMTIHON NAZORATI (MONITOR)
# ==========================================
@router.callback_query(F.data.startswith("live_monitor_"))
async def cb_live_monitor(callback: CallbackQuery):
    """Imtihon jarayonini jonli kuzatish oynasi (yangilash, majburiy yakunlash)"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz:
            await callback.answer("Test topilmadi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!", show_alert=True)
            return

        # Studentlar va urinishlar ma'lumotlarini olish
        students = await session.scalars(
            select(Student).where(Student.quiz_id == quiz_id).order_by(Student.full_name)
        )
        students = students.all()

        # Urinishlar
        attempts_stmt = select(StudentAttempt).where(
            StudentAttempt.student_id.in_([s.id for s in students])
        )
        attempts_res = await session.execute(attempts_stmt)
        attempts = attempts_res.scalars().all()
        attempts_map = {a.student_id: a for a in attempts}

        # Qolgan vaqtni hisoblash
        now = local_now()
        remaining_minutes = 0
        if quiz.end_time:
            delta = quiz.end_time - now
            remaining_minutes = max(0, int(delta.total_seconds() / 60))

        completed_list = []
        solving_list = []
        idle_list = []

        for s in students:
            att = attempts_map.get(s.id)
            if att:
                # O'quvchi yechish jarayonida javob bergan savollar soni
                ans_count_stmt = select(func.count(StudentAnswer.id)).where(
                    StudentAnswer.attempt_id == att.id, StudentAnswer.selected_option_index.isnot(None)
                )
                ans_count_res = await session.execute(ans_count_stmt)
                ans_count = ans_count_res.scalar()

                if att.completed_at:
                    completed_list.append(f"• {s.full_name} ({att.score} ball, tugatgan)")
                else:
                    solving_list.append(f"• {s.full_name} (Progress: {ans_count} ta yechdi)")
            else:
                if s.telegram_id:
                    idle_list.append(f"• {s.full_name} (Kutish zalida, hali start bosmadi)")
                else:
                    idle_list.append(f"• {s.full_name} (Tizimga kirmagan)")

        monitor_text = (
            f"📊 **IMTIHON JONLI NAZORATI**\n"
            f"📚 Test: *{quiz.title}* | Kod: `{quiz.code}`\n"
            f"⏳ Vaqt tugashiga: *{remaining_minutes} daqiqa* qoldi\n\n"
            f"✅ **Tugatganlar ({len(completed_list)}):**\n"
            + "\n".join(completed_list)
            + "\n\n"
            f"⏳ **Yechayotganlar ({len(solving_list)}):**\n"
            + "\n".join(solving_list)
            + "\n\n"
            f"❌ **Hali boshlamaganlar ({len(idle_list)}):**\n"
            + "\n".join(idle_list)
        )

        from aiogram.types import InlineKeyboardMarkup, InlineKeyboardButton

        buttons = [
            [
                InlineKeyboardButton(text="🔄 Yangilash", callback_data=f"live_monitor_{quiz_id}"),
                InlineKeyboardButton(
                    text="🛑 Majburiy Yakunlash", callback_data=f"terminate_quiz_{quiz_id}"
                ),
            ],
            [InlineKeyboardButton(text="🔙 Orqaga", callback_data=f"manage_quiz_{quiz_id}")],
        ]

        await callback.message.edit_text(
            monitor_text,
            reply_markup=InlineKeyboardMarkup(inline_keyboard=buttons),
            parse_mode="Markdown",
        )
        await callback.answer()


@router.callback_query(F.data.startswith("terminate_quiz_"))
async def cb_terminate_quiz(callback: CallbackQuery, bot: Bot):
    """Imtihonni vaqtidan oldin majburiy yakunlash (Tugatish)"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz or quiz.status != "active":
            await callback.answer("Imtihonni yakunlab bo'lmaydi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni boshqara olasiz!", show_alert=True)
            return

        # Holatni completed qilish
        quiz.status = "completed"
        quiz.end_time = local_now()

        # Imtihon topshirayotgan barcha faol studentlarni yakunlash va auto-logout qilish
        students = await session.scalars(
            select(Student).where(Student.quiz_id == quiz_id)
        )
        students = students.all()

        active_student_tg_ids = []
        for s in students:
            if s.telegram_id:
                active_student_tg_ids.append(s.telegram_id)

            # Urinishni to'xtatish
            att_stmt = select(StudentAttempt).where(
                StudentAttempt.student_id == s.id, StudentAttempt.completed_at.is_(None)
            )
            att_res = await session.execute(att_stmt)
            attempt = att_res.scalar_one_or_none()

            if attempt:
                attempt.completed_at = quiz.end_time
                # Javoblarni baholash
                answers_stmt = select(StudentAnswer).where(StudentAnswer.attempt_id == attempt.id)
                answers_res = await session.execute(answers_stmt)
                answers = answers_res.scalars().all()

                correct_count = sum(1 for a in answers if a.is_correct)
                attempt.score = correct_count
                attempt.total_questions = len(answers)

            # Auto logout
            s.telegram_id = None

        await session.commit()

    # Barcha topshirayotgan studentlarga test yopilganini bildirish (ID larni oldindan saqlab qo'ydik)
    for tg_id in active_student_tg_ids:
        try:
            await bot.send_message(
                tg_id,
                "🛑 **Imtihon Moderator tomonidan yakunlandi!**\nTizimdan avtomat chiqdingiz (Logout). Natijangiz saqlandi.",
            )
        except Exception:
            pass

    await callback.message.edit_text(
        "🛑 **Imtihon yakunlandi.** Barcha tizimdan chiqarildi va natijalar saqlandi.",
        parse_mode="Markdown"
    )
    # Avtomatik ravishda asosiy ro'yxatni ko'rsatish
    await show_active_quizzes(callback.message, state)


# ==========================================
# 9. EXCEL NATIJALARINI YUKLASH (EXPORTER)
# ==========================================
@router.callback_query(F.data.startswith("download_results_"))
async def cb_download_results(callback: CallbackQuery):
    """Baholash matritsasi va umumiy natijalar Excel faylini generatsiya qilib Moderatorga jo'natish"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz:
            await callback.answer("Test topilmadi!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarning natijalarini yuklab ola olasiz!", show_alert=True)
            return

        # Savollarni olish (ID bo'yicha tartiblash uchun)
        questions_stmt = select(Question).where(Question.quiz_id == quiz_id).order_by(Question.id)
        questions_res = await session.execute(questions_stmt)
        questions = questions_res.scalars().all()
        q_ids = [q.id for q in questions]
        q_count = len(q_ids)

        # Studentlarni olish
        students_stmt = select(Student).where(Student.quiz_id == quiz_id).order_by(Student.full_name)
        students_res = await session.execute(students_stmt)
        students = students_res.scalars().all()
        students_list = [{"id": s.id, "full_name": s.full_name} for s in students]

        # Urinishlar va javoblar
        attempts_stmt = select(StudentAttempt).where(
            StudentAttempt.student_id.in_([s.id for s in students])
        )
        attempts_res = await session.execute(attempts_stmt)
        attempts = attempts_res.scalars().all()

        attempts_map = {}
        for att in attempts:
            answers_stmt = select(StudentAnswer).where(StudentAnswer.attempt_id == att.id)
            answers_res = await session.execute(answers_stmt)
            answers = answers_res.scalars().all()

            # Savol id si bo'yicha javoblarning to'g'ri-xatoligi
            # Excel matritsasida savol tartibi (1, 2, ... m) bo'lishi kerak.
            # Shuning uchun original question_id ning q_ids ro'yxatidagi 1-indeksli o'rnini xaritaymiz.
            answers_dict = {}
            for ans in answers:
                try:
                    q_idx = q_ids.index(ans.question_id) + 1  # 1 dan m gacha tartib
                    answers_dict[q_idx] = ans.is_correct
                except ValueError:
                    pass

            attempts_map[att.student_id] = {
                "score": att.score,
                "completed": att.completed_at is not None,
                "answers": answers_dict,
            }

    # Excel hosil qilish
    excel_bytes = generate_results_excel(quiz.title, students_list, q_count, attempts_map)

    await callback.message.answer_document(
        BufferedInputFile(excel_bytes, filename=f"natijalar_{quiz.code}.xlsx"),
        caption=f"📊 Test: {quiz.title}\nKodi: {quiz.code}\n\nBaholash matritsasi va umumiy natijalar Excel hisoboti.",
    )
    await callback.answer()


# ==========================================
# 10. TESTNI O'CHIRISH
# ==========================================
@router.callback_query(F.data.startswith("delete_quiz_"))
async def cb_delete_quiz(callback: CallbackQuery):
    """Testni ma'lumotlar bazasidan butunlay o'chirish (savollari va rasmlari bilan kaskad o'chadi)"""
    quiz_id = int(callback.data.split("_")[2])

    async with async_session() as session:
        quiz = await session.get(Quiz, quiz_id)
        if not quiz:
            await callback.answer("Test allaqachon o'chirilgan!", show_alert=True)
            return

        if quiz.teacher_id != callback.from_user.id and callback.from_user.id not in ADMIN_IDS:
            await callback.answer("Siz faqat o'zingiz yaratgan testlarni o'chira olasiz!", show_alert=True)
            return

        title = quiz.title
        await session.delete(quiz)
        await session.commit()

    await callback.message.answer(f"✅ *{title}* testi va unga tegishli barcha urinishlar bazadan butunlay o'chirildi.", parse_mode="Markdown")
    await callback.message.delete()
    await callback.answer()

@router.message(F.text == "Tariflar va Imkoniyatlar")
async def show_tariffs_guide(message: Message):
    text = (
        "💎 <b>TARIFLAR VA IMKONIYATLAR</b>\n\n"
        "Hurmatli ustoz, botimiz orqali test olish jarayonini avtomatlashtirish endi juda qulay!\n\n"
        "🎁 <b>BEPUL DEMO REJIMI</b> (Barchaga avtomatik beriladi)\n"
        "• <b>Testlar soni:</b> faqat 1 marta\n"
        "• <b>O'quvchilar sig'imi:</b> 50 nafargacha\n"
        "• <b>Qo'shimcha bepul limit:</b> 5 ta hamkasb taklif qilib, yana bepul limit olish mumkin!\n\n"
        
        "🎟 <b>1 MARTALIK LIMIT</b>\n"
        "• <b>Maqsad:</b> doimiy Premium olmasdan, vaqti-vaqti bilan bittadan test o'tkazmoqchi bo'lganlar uchun.\n"
        "• <b>Testlar soni:</b> har bir limit 1 ta test jarayoni\n"
        "• <b>O'quvchilar sig'imi:</b> 50 nafargacha\n\n"
        
        "🌟 <b>PREMIUM OBUNA (1 OYLIK)</b>\n"
        "• <b>Testlar soni:</b> 1 oy davomida CHEKSIZ ♾️\n"
        "• <b>O'quvchilar sig'imi:</b> 50 nafargacha (bitta test uchun)\n"
        "• <b>Foydasi:</b> oy davomida xohlagancha test olish imkoniyati!\n\n"
        
        "<i>💡 Tariflar narxlari bilan tanishish hamda o'zingizga mos tarifni xarid qilish uchun darhol <b>\"Adminga murojaat\"</b> tugmasini bosing!</i>"
    )
    
    kb = InlineKeyboardMarkup(inline_keyboard=[
        [InlineKeyboardButton(text="🤝 Taklif havolasini olish (Bepul limit)", callback_data="get_ref_link")]
    ])
    await message.answer(text, parse_mode="HTML", reply_markup=kb)


@router.callback_query(F.data == "get_ref_link")
async def cb_get_ref_link(callback: CallbackQuery):
    user_id = callback.from_user.id
    bot_info = await callback.bot.get_me()
    ref_link = f"https://t.me/{bot_info.username}?start=ref_{user_id}"
    
    async with async_session() as session:
        user_stmt = select(User).where(User.id == user_id)
        db_u = (await session.execute(user_stmt)).scalar_one_or_none()
        ref_count = db_u.referral_count if db_u else 0

    text = (
        f"🤝 **Taklif tizimi orqali bepul limit oling!**\n\n"
        f"Quyidagi taklif havolasini hamkasb ustozlarga yuboring. Ular botga kirib `/start` bosishlari bilanoq sizning hisobingizga 1 ta taklif qo'shiladi:\n\n"
        f"🔗 Taklif havolangiz: `{ref_link}`\n\n"
        f"📊 Hozirgi ko'rsatkich: *{ref_count}/5* hamkasb taklif qilingan.\n"
        f"🎁 5 taga yetganda **1 ta bepul test yaratish limiti** (onetime) olasiz!"
    )
    await callback.message.answer(text, parse_mode="Markdown")
    await callback.answer()
